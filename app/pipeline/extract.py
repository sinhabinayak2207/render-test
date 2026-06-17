"""Multi-format document extraction -> markdown.

Chain (accuracy-first):
  PDF  -> per page: PyMuPDF text if a text layer exists, else OCR the rendered
          page image (PaddleOCR; fall back to gpt-4o-mini vision when Paddle is
          missing / low-confidence / errors).
  HTML -> BeautifulSoup + markdownify
  XLS/XLSX -> markdown tables
  DOCX -> python-docx ; DOC -> LibreOffice (when available)
  JSON -> passthrough

Pure-Python rasterization via PyMuPDF (no poppler needed).
"""
from __future__ import annotations

import hashlib
import io
import logging
import os
import subprocess
import tempfile
from dataclasses import dataclass, field

from ..config import settings
from ..llm_extract import vision_ocr

log = logging.getLogger("extract")

# ── PaddleOCR lazy singleton ─────────────────────────────────────────────────
_PADDLE = None
_PADDLE_TRIED = False


def _paddle():
    global _PADDLE, _PADDLE_TRIED
    if _PADDLE_TRIED:
        return _PADDLE
    _PADDLE_TRIED = True
    try:
        from paddleocr import PaddleOCR

        last = None
        # Constructor differs across versions (3.x dropped show_log/use_angle_cls).
        # Disable the heavy doc-orientation/unwarping models — useless for tender
        # scans and the main cause of slow per-page inference.
        for kwargs in (
            {"lang": settings.ocr_lang, "use_doc_orientation_classify": False,
             "use_doc_unwarping": False, "use_textline_orientation": False},
            {"lang": settings.ocr_lang},
            {"use_angle_cls": True, "lang": settings.ocr_lang, "show_log": False},
        ):
            try:
                _PADDLE = PaddleOCR(**kwargs)
                break
            except Exception as e:  # noqa: BLE001 — try the next signature
                last, _PADDLE = e, None
        if _PADDLE is None:
            raise last or RuntimeError("PaddleOCR init failed")
        log.info("PaddleOCR ready")
    except Exception as exc:  # noqa: BLE001
        log.warning("PaddleOCR unavailable (%s) — using gpt-4o-mini vision OCR fallback", exc)
        _PADDLE = None
    return _PADDLE


_PADDLE_DUMPED = False


def _parse_paddle(res) -> tuple[list[str], list[float]]:
    """Parse PaddleOCR output across versions.

    3.x `predict()` returns a single OCRResult (dict subclass) OR a list of them,
    each with `rec_texts`/`rec_scores`. 2.x `ocr()` returns [[ [box,(text,conf)], ... ]].
    """
    global _PADDLE_DUMPED
    lines: list[str] = []
    confs: list[float] = []
    if res is None:
        return lines, confs
    pages = res if isinstance(res, list) else [res]  # normalize single result -> list
    for page in pages:
        rec_texts = rec_scores = None
        if isinstance(page, dict):
            rec_texts, rec_scores = page.get("rec_texts"), page.get("rec_scores")
            if rec_texts is None:  # some builds nest under 'res'/'result'
                inner = page.get("res") or page.get("result") or {}
                if isinstance(inner, dict):
                    rec_texts, rec_scores = inner.get("rec_texts"), inner.get("rec_scores")
        else:
            rec_texts = getattr(page, "rec_texts", None)
            rec_scores = getattr(page, "rec_scores", None)
        if rec_texts:  # PaddleOCR 3.x
            lines.extend(str(t) for t in rec_texts)
            confs.extend(float(s) for s in (rec_scores or []))
            continue
        if isinstance(page, (list, tuple)):  # PaddleOCR 2.x
            for item in page:
                try:
                    lines.append(item[1][0])
                    confs.append(float(item[1][1]))
                except Exception:  # noqa: BLE001
                    continue
    if not lines and not _PADDLE_DUMPED:  # one-shot: reveal the shape if still unparsed
        _PADDLE_DUMPED = True
        keys = list(pages[0].keys()) if (pages and isinstance(pages[0], dict)) else "n/a"
        log.warning("PaddleOCR result unparsed: type=%s keys=%s sample=%s",
                    type(res).__name__, keys, repr(res)[:400])
    return lines, confs


@dataclass
class ExtractResult:
    fmt: str
    markdown: str
    ocr_used: bool = False
    content_hash: str = ""
    meta: dict = field(default_factory=dict)
    pages: list = field(default_factory=list)  # per-page text (for page-level source)


def _hash(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


# ── OCR a single rendered page ───────────────────────────────────────────────
def _ocr_png(png: bytes) -> str:
    paddle = _paddle()
    text, conf = "", 0.0
    if paddle is not None:
        try:
            import numpy as np
            from PIL import Image

            img = np.array(Image.open(io.BytesIO(png)).convert("RGB"))
            raw = None
            for caller in (
                lambda: paddle.predict(img),        # PaddleOCR 3.x (returns a generator!)
                lambda: paddle.ocr(img),
                lambda: paddle.ocr(img, cls=True),  # PaddleOCR 2.x
            ):
                try:
                    raw = caller()
                    break
                except Exception:  # noqa: BLE001 — try next call style
                    continue
            res = list(raw) if raw is not None else []  # force-consume generator
            lines, confs = _parse_paddle(res)
            text = "\n".join(lines).strip()
            conf = sum(confs) / len(confs) if confs else 0.0
            log.info("PaddleOCR page → %d chars, conf=%.2f", len(text), conf)
        except Exception as exc:  # noqa: BLE001
            log.warning("PaddleOCR failed on page (%s) — trying gpt-4o-mini vision", exc)
            text, conf = "", 0.0
    # Fallback to gpt-4o-mini vision when Paddle is absent / weak / empty.
    if settings.enable_vision_fallback and (not text or conf < settings.ocr_min_confidence):
        log.info("→ gpt-4o-mini vision OCR fallback (paddle text=%d chars, conf=%.2f)", len(text), conf)
        try:
            g = vision_ocr(png)
            if len(g) > len(text):
                return g
        except Exception as exc:  # noqa: BLE001
            log.warning("gpt-4o-mini vision OCR failed: %s", exc)
    return text


# ── format handlers ──────────────────────────────────────────────────────────
def _pdf(content: bytes) -> ExtractResult:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=content, filetype="pdf")
    pages, ocr_used = [], False
    thr = settings.pdf_text_min_chars_per_page
    for page in doc:
        ptext = page.get_text().strip()
        if len(ptext) >= thr:
            pages.append(ptext)
        elif settings.enable_ocr:
            otext = _ocr_png(page.get_pixmap(matrix=fitz.Matrix(2, 2)).tobytes("png"))
            if otext:
                ocr_used = True
            pages.append(otext)
        else:
            pages.append("")  # scanned page, OCR disabled (text-first mode)
    md = "\n\n".join(f"<!-- page {i} -->\n{t}" for i, t in enumerate(pages, 1)).strip()
    return ExtractResult("scanned_pdf" if ocr_used else "digital_pdf", md, ocr_used, pages=pages)


def _html(content: bytes) -> ExtractResult:
    from bs4 import BeautifulSoup
    from markdownify import markdownify as mdify

    soup = BeautifulSoup(content, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    md = mdify(str(soup)).strip()
    plain = soup.get_text("\n", strip=True)  # the "Ctrl+A" raw text — never misses content
    # Keep markdown (tables/links) unless it lost too much vs the raw selection.
    best = md if len(md) >= len(plain) * 0.6 else plain
    return ExtractResult("html", best, pages=[best])


def _xlsx(content: bytes) -> ExtractResult:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"### Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                out.append("| " + " | ".join(cells) + " |")
    return ExtractResult("xlsx", "\n".join(out).strip())


def _xls(content: bytes) -> ExtractResult:
    import xlrd

    book = xlrd.open_workbook(file_contents=content)
    out = []
    for sh in book.sheets():
        out.append(f"### Sheet: {sh.name}")
        for r in range(sh.nrows):
            cells = [str(sh.cell_value(r, c)) for c in range(sh.ncols)]
            if any(x.strip() for x in cells):
                out.append("| " + " | ".join(cells) + " |")
    return ExtractResult("xls", "\n".join(out).strip())


def _docx(content: bytes) -> ExtractResult:
    import docx

    d = docx.Document(io.BytesIO(content))
    out = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            out.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")
    return ExtractResult("docx", "\n".join(out).strip())


def _doc(content: bytes) -> ExtractResult:
    """Legacy binary .doc via LibreOffice headless (if installed)."""
    soffice = next((p for p in ("soffice", "libreoffice") if _which(p)), None)
    if not soffice:
        return ExtractResult("doc", "[.doc not extracted — LibreOffice not available on this host]")
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "in.doc")
        with open(src, "wb") as f:
            f.write(content)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmp, src],
                check=True, capture_output=True, timeout=120,
            )
            txt = os.path.join(tmp, "in.txt")
            if os.path.exists(txt):
                return ExtractResult("doc", open(txt, encoding="utf-8", errors="ignore").read().strip())
        except Exception as exc:  # noqa: BLE001
            log.warning(".doc conversion failed: %s", exc)
    return ExtractResult("doc", "[.doc conversion failed]")


def _which(cmd: str) -> str | None:
    from shutil import which

    return which(cmd)


_IMAGE_EXTS = {"png", "jpg", "jpeg", "tif", "tiff", "bmp", "webp", "gif"}


def _to_png(content: bytes) -> bytes | None:
    try:
        import io

        from PIL import Image

        im = Image.open(io.BytesIO(content)).convert("RGB")
        buf = io.BytesIO()
        im.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:  # noqa: BLE001
        return None


def _image(content: bytes) -> ExtractResult:
    if not settings.enable_vision_fallback:
        return ExtractResult("image", "")
    png = _to_png(content)
    text = vision_ocr(png) if (settings.enable_vision_fallback and png) else ""
    return ExtractResult("image", text, ocr_used=True)


def vision_recover(name: str, content: bytes) -> str:
    """Last-resort reader: render a document to image(s) and OCR via gpt-4o-mini vision.

    Used when normal parsing returns empty/errors. Works for PDFs (rasterized
    per page) and image files; returns "" for formats with no image form (.doc/.xls).
    """
    if not settings.enable_vision_fallback:
        return ""
    ext = os.path.splitext((name or "").lower())[1].lstrip(".")
    try:
        if ext == "pdf":
            import fitz

            doc = fitz.open(stream=content, filetype="pdf")
            parts = []
            for page in doc:
                png = page.get_pixmap(matrix=fitz.Matrix(2, 2)).tobytes("png")
                parts.append(vision_ocr(png))
            return "\n\n".join(p for p in parts if p).strip()
        png = _to_png(content)
        if png:
            return vision_ocr(png)
    except Exception as exc:  # noqa: BLE001
        log.warning("vision_recover failed for %s: %s", name, exc)
    return ""


def _json(content: bytes) -> ExtractResult:
    import json as _j

    try:
        obj = _j.loads(content.decode("utf-8", "ignore"))
        md = "```json\n" + _j.dumps(obj, indent=2, ensure_ascii=False)[:20000] + "\n```"
        return ExtractResult("json", md, meta={"parsed": obj})
    except Exception:  # noqa: BLE001
        return ExtractResult("json", content.decode("utf-8", "ignore")[:20000])


# ── dispatch ──────────────────────────────────────────────────────────────────
def extract(name: str, content: bytes) -> ExtractResult:
    ext = os.path.splitext((name or "").lower())[1].lstrip(".")
    try:
        if ext == "pdf":
            res = _pdf(content)
        elif ext in ("html", "htm"):
            res = _html(content)
        elif ext == "xlsx":
            res = _xlsx(content)
        elif ext == "xls":
            res = _xls(content)
        elif ext == "docx":
            res = _docx(content)
        elif ext == "doc":
            res = _doc(content)
        elif ext == "json":
            res = _json(content)
        elif ext in _IMAGE_EXTS:
            res = _image(content)
        else:
            # last resort: try plain text
            res = ExtractResult(ext or "unknown", content.decode("utf-8", "ignore").strip()[:20000])
    except Exception as exc:  # noqa: BLE001
        log.exception("extraction failed for %s", name)
        res = ExtractResult(ext or "unknown", f"[extraction error: {exc}]")
    res.content_hash = _hash(content)
    return res
