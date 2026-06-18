"""Build the Tender Intelligence Report as a DOCX from Supabase rows.

Reads a run's tenders (default = latest run) + the company profile, and assembles
a Word document: matching profile -> executive summary -> per-tender deep dives
(eligible/partial) -> rejected. Margin (eligibility_flags) get highlighted.

Run:  python -m app.report_docx [run_id]   (omit run_id for the latest run)
"""
from __future__ import annotations

import json
import sys
import urllib.request

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .config import settings

_H = {"apikey": settings.supabase_service_key, "Authorization": f"Bearer {settings.supabase_service_key}"}
RED = RGBColor(0xB4, 0x23, 0x18)
AMBER = RGBColor(0x9A, 0x67, 0x00)
GREEN = RGBColor(0x1A, 0x7F, 0x37)
NAVY = RGBColor(0x08, 0x26, 0x3F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "08263F"
LITE_HEX = "EAF0F6"
VORDER = {"ELIGIBLE": 0, "PARTIAL": 1, "INELIGIBLE": 2, "EXCLUDED": 3, "PENDING": 4}
VCOLOR = {"ELIGIBLE": GREEN, "PARTIAL": AMBER, "INELIGIBLE": RED, "EXCLUDED": RED, "PENDING": NAVY}
VLABEL = {"ELIGIBLE": "ELIGIBLE", "PARTIAL": "PARTIALLY ELIGIBLE",
          "INELIGIBLE": "REJECTED", "EXCLUDED": "REJECTED (out of scope)", "PENDING": "PENDING"}


def _vlabel(v):
    return VLABEL.get(v, v or "")
_EMPTY = {"", "[]", "{}", "null", "none", "n/a", "na", "-", "—", "[ ]"}


def _get(path: str):
    return json.load(urllib.request.urlopen(urllib.request.Request(f"{settings.supabase_url}/rest/v1/{path}", headers=_H), timeout=60))


def _clean(v):
    """Drop empty-ish values, incl. the string artefacts the LLM sometimes returns."""
    if v in (None, [], {}):
        return None
    if isinstance(v, str) and v.strip().lower() in _EMPTY:
        return None
    return v


def _money_cr(inr) -> str:
    if not inr:
        return "Value not disclosed"
    cr = inr / 1e7
    return f"Rs. {cr:.2f} Cr" if cr >= 0.01 else "Value not disclosed"


def _clause(c) -> str:
    if isinstance(c, dict):
        t = c.get("text") or ""
        pg = c.get("pages")
        return t + (f"  [pp {pg}]" if pg else "")
    return str(c)


def _shade(cell, hex_fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell(cell, text, *, bold=False, color=None, fill=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if fill:
        _shade(cell, fill)


def _header_row(table, labels):
    for i, lab in enumerate(labels):
        _cell(table.rows[0].cells[i], lab, bold=True, color=WHITE, fill=NAVY_HEX)


def _kv(doc, pairs):
    pairs = [(k, _clean(v)) for k, v in pairs]
    pairs = [(k, v) for k, v in pairs if v is not None]
    if not pairs:
        return
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in pairs:
        r = t.add_row().cells
        _cell(r[0], k, bold=True, fill=LITE_HEX)
        r[1].text = str(v)


def _bullets(doc, items, fmt=str):
    for it in (items or []):
        if _clean(it) is None:
            continue
        doc.add_paragraph(fmt(it), style="List Bullet")


def _sub(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = NAVY


def _highlight(doc, text, color=AMBER):
    p = doc.add_paragraph()
    r = p.add_run("⚠ " + text)  # warning sign
    r.bold = True
    r.font.color.rgb = color


def _real_docs(t):
    """Actual tender files (PDF / xls / doc) — drop the saved HTML detail page."""
    out = []
    for d in (t.get("downloaded_docs") or []):
        fmt = str(d.get("format") or "").lower()
        if fmt == "html" or str(d.get("kind") or "").lower() == "html":
            continue
        out.append(d)
    return out


# ── per-tender deep dive ──────────────────────────────────────────────────────
def _deep_block(doc, t, idx):
    ed = t.get("extracted_data") or {}
    extras = (ed.get("extras") or {}) if isinstance(ed, dict) else {}
    h = doc.add_heading(level=2)
    r = h.add_run(f"#{idx}   [{_vlabel(t.get('verdict'))}]   {t.get('title') or ''}")
    r.font.color.rgb = VCOLOR.get(t.get("verdict"), NAVY)

    dur = t.get("project_duration")
    dur = f"{dur} days" if str(dur).strip().isdigit() else dur

    _sub(doc, "1. Basic Tender Information")
    _kv(doc, [
        ("Issuing Authority", t.get("issuing_authority")),
        ("Authority Location", t.get("issuing_authority_location")),
        ("Authority Contact", t.get("authority_contact")),
        ("Portal", t.get("portal_name")),
        ("Reference No.", t.get("reference_number")),
        ("Estimated Value", _money_cr(t.get("estimated_value"))),
        ("EMD", _money_cr(t.get("emd_amount"))),
        ("Published", (t.get("published_at") or "")[:10] or None),
        ("Submission Deadline", t.get("closing_date")),
        ("Opening", t.get("opening_date")),
        ("Tender Type", (f"{t.get('tender_type')}" + (f" (conf: {ed.get('tender_type_confidence')})" if ed.get('tender_type_confidence') else "")) if t.get("tender_type") else None),
        ("Procurement Model", t.get("procurement_model")),
        ("Commercial Structure", t.get("commercial_model")),
    ])

    _sub(doc, "2. Bid Decision Summary")
    _kv(doc, [
        ("Strategic Fit Score", f"{(t.get('competitiveness_score') or 0)/10:.1f} / 10"),
        ("Basis", t.get("strategic_fit_basis")),
        ("Recommendation", t.get("narrative_fit")),
        ("Key Business Insight", t.get("key_business_insight")),
    ])

    _sub(doc, "3. Scope of Work")
    _kv(doc, [("Scope Summary", t.get("scope_summary")),
              ("Project Duration", dur),
              ("Location", t.get("location_of_execution")),
              ("SoW Page Refs", t.get("sow_page_refs"))])
    if _clean(t.get("key_deliverables")):
        doc.add_paragraph("Key Deliverables:").runs[0].bold = True
        _bullets(doc, t.get("key_deliverables"))

    _sub(doc, "4. Risk Assessment")
    _kv(doc, [("Compliance Complexity", extras.get("compliance_complexity")),
              ("Compliance Basis", t.get("compliance_basis")),
              ("Overall Risk", t.get("risk_level")),
              ("Risk (plain English)", t.get("risk_layperson_explanation"))])
    if _clean(t.get("unusual_clauses")):
        doc.add_paragraph("Unusual / Risky Clauses:").runs[0].bold = True
        _bullets(doc, t.get("unusual_clauses"), _clause)
    if _clean(t.get("penalty_clauses")):
        doc.add_paragraph("Penalty Clauses:").runs[0].bold = True
        _bullets(doc, t.get("penalty_clauses"), _clause)

    _sub(doc, "5. Eligibility Check (CS Direkt Rubric)")
    rows = t.get("eligibility_check") or []
    if rows:
        tb = doc.add_table(rows=1, cols=3)
        tb.style = "Table Grid"
        _header_row(tb, ["Dimension", "Status", "Detail"])
        for row in rows:
            c = tb.add_row().cells
            c[0].text = str(row.get("dimension", ""))
            c[1].text = str(row.get("status", ""))
            c[2].text = str(row.get("detail", ""))
    for fl in (t.get("eligibility_flags") or []):
        band = fl.get("band")
        col = AMBER if band == "borderline_5" else RED
        verb = "within ±5% (borderline — still eligible)" if band == "borderline_5" else f"{fl.get('over_pct')}% over capacity"
        _highlight(doc, f"{fl.get('field','').replace('_',' ').title()}: required Rs.{fl.get('required')} Cr "
                        f"vs capacity Rs.{fl.get('capacity')} Cr — {verb}", col)

    if _clean(t.get("gaps_to_address")):
        _sub(doc, "6. Gaps to Address")
        _bullets(doc, t.get("gaps_to_address"))

    if _clean(t.get("pre_bid_queries")):
        _sub(doc, "7. Pre-Bid Queries")
        for q in t.get("pre_bid_queries"):
            if isinstance(q, dict):
                # New consultant shape (matches report_pdf._prebid / narrative.py):
                # priority / clause_reference / existing_requirement / observation /
                # question / strategic_objective / expected_benefit.
                pri = q.get("priority") or ""
                ref = q.get("clause_reference") or ""
                head = " · ".join(x for x in (pri, ref) if x)
                p = doc.add_paragraph()
                p.add_run(f"[{head}] " if head else "").bold = True
                p.add_run(q.get("question", ""))
                for lbl, key in (("Current requirement", "existing_requirement"),
                                 ("Observation", "observation"),
                                 ("Objective", "strategic_objective"),
                                 ("Expected benefit", "expected_benefit")):
                    if q.get(key):
                        rp = doc.add_paragraph()
                        rp.add_run(f"{lbl}: {q[key]}").italic = True

    if extras.get("pricing_feasibility") or extras.get("epc_estimate_cr"):
        _sub(doc, "8. L1 / Commercial Analysis")
        _kv(doc, [("Pricing Feasibility", extras.get("pricing_feasibility")),
                  ("EPC Estimate", f"Rs. {extras.get('epc_estimate_cr')} Cr" if extras.get("epc_estimate_cr") else None)])

    docs = _real_docs(t)
    if docs:
        _sub(doc, "9. Source Documents (PDF / BOQ)")
        for d in docs:
            url = d.get("storage_url") or d.get("public_url") or ""
            label = f"{d.get('kind') or 'Document'} — {d.get('name')}"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(label + ": ").bold = True
            p.add_run(url)
    doc.add_paragraph("")


def _rejected_block(doc, t, idx):
    h = doc.add_heading(level=2)
    r = h.add_run(f"#{idx}   [{_vlabel(t.get('verdict'))}]   {t.get('title') or ''}")
    r.font.color.rgb = RED
    _kv(doc, [("Issuing Authority", t.get("issuing_authority")),
              ("Portal", t.get("portal_name")),
              ("Reference No.", t.get("reference_number")),
              ("Estimated Value", _money_cr(t.get("estimated_value"))),
              ("Matched Keyword", t.get("matched_keyword")),
              ("Excluded Reason", t.get("excluded_reason"))])
    if _clean(t.get("disqualification_triggers")):
        doc.add_paragraph("Disqualification Triggers:").runs[0].bold = True
        for d in t["disqualification_triggers"]:
            if isinstance(d, dict):
                doc.add_paragraph(f"{d.get('name','')}: {d.get('evidence','')}", style="List Bullet")
    if _clean(t.get("business_logic_explanation")):
        doc.add_paragraph(t["business_logic_explanation"])
    elif _clean(t.get("reasons_rejected")):
        _bullets(doc, t.get("reasons_rejected"))
    doc.add_paragraph("")


# ── main build ────────────────────────────────────────────────────────────────
def build_report(run_id: str | None = None, out_path: str = "/mnt/c/Users/sinha/Downloads/CSD_tender_report.docx") -> str:
    if run_id:
        tenders = _get(f"tenders?run_id=eq.{run_id}&select=*")
        scope_label = f"Run {run_id[:8]}"
    else:
        # Universal default: report EVERY tender currently in the table (all runs).
        tenders = _get("tenders?select=*&order=created_at.desc&limit=1000")
        scope_label = "All tenders"
    prof = (_get("company_profiles?user_id=is.null&select=*") or [{}])[0]
    tenders = [t for t in tenders if t.get("verdict") != "EXCLUDED"]  # excluded → not in the report at all
    tenders.sort(key=lambda t: (VORDER.get(t.get("verdict"), 9), -(t.get("competitiveness_score") or 0)))

    counts = {k: sum(1 for t in tenders if t.get("verdict") == k) for k in ("ELIGIBLE", "PARTIAL", "INELIGIBLE")}
    doc = Document()
    ttl = doc.add_heading(level=0).add_run("TENDER INTELLIGENCE REPORT")
    ttl.font.color.rgb = NAVY
    doc.add_paragraph(f"CS Direkt — Tender Search & Eligibility Analysis   ·   {scope_label}   ·   {len(tenders)} tenders")
    kp = doc.add_paragraph()
    for lab, key, col in [("Eligible", "ELIGIBLE", GREEN), ("Partially Eligible", "PARTIAL", AMBER), ("Rejected", "INELIGIBLE", RED)]:
        rr = kp.add_run(f"   {lab}: {counts[key]}    ")
        rr.bold = True
        rr.font.color.rgb = col

    _sub(doc, "Matching Profile (eligibility & preferences)")
    _kv(doc, [
        ("Company", prof.get("company_name")),
        ("Min / Max tender value (Cr)", f"{prof.get('min_tender_value_cr')} / {prof.get('max_tender_value_cr') or 'no limit'}"),
        ("Turnover capacity (Cr)", prof.get("turnover_latest_cr")),
        ("Net worth (Cr)", prof.get("net_worth_latest_cr")),
        ("Eligible / Partial score cutoff", f"{prof.get('eligible_min_score')} / {prof.get('partial_min_score')}"),
        ("Service lines", ", ".join(prof.get("service_lines") or []) or None),
        ("Exclude keywords", (", ".join((prof.get("exclude_keywords") or [])[:12]) + " …") if prof.get("exclude_keywords") else None),
    ])

    doc.add_heading("Executive Summary", level=1)
    tb = doc.add_table(rows=1, cols=6)
    tb.style = "Table Grid"
    _header_row(tb, ["Sl", "Tender Title", "Issuing Authority", "Estimated Value", "Strategic Fit Status", "Key Business Action / Insight"])
    for n, t in enumerate(tenders, 1):
        c = tb.add_row().cells
        c[0].text = str(n)
        c[1].text = (t.get("title") or "")[:90]
        c[2].text = (t.get("issuing_authority") or "")[:40]
        c[3].text = _money_cr(t.get("estimated_value"))
        _cell(c[4], _vlabel(t.get("verdict")), bold=True, color=VCOLOR.get(t.get("verdict"), NAVY))
        c[5].text = (t.get("key_business_insight") or t.get("excluded_reason") or "")[:160]

    elig = [t for t in tenders if t.get("verdict") == "ELIGIBLE"]
    part = [t for t in tenders if t.get("verdict") == "PARTIAL"]
    rej = [t for t in tenders if t.get("verdict") == "INELIGIBLE"]
    if elig:
        doc.add_page_break(); doc.add_heading("Section A — Fully Eligible", level=1)
        for i, t in enumerate(elig, 1): _deep_block(doc, t, i)
    if part:
        doc.add_heading("Section B — Partially Eligible", level=1)
        for i, t in enumerate(part, 1): _deep_block(doc, t, i)
    if rej:
        doc.add_heading("Section C — Rejected", level=1)
        for i, t in enumerate(rej, 1): _rejected_block(doc, t, i)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    rid = sys.argv[1] if len(sys.argv) > 1 else None
    print("DOCX written:", build_report(rid))
